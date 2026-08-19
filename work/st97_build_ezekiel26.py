from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


OUT = Path("/Users/yoshiakinagumo/Documents/Obsidian Vault/outputs/エゼキエル26章_早天メッセージ資料.docx")
FONT = "Noto Sans JP"
BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(32, 32, 32)
MUTED = RGBColor(90, 90, 90)


def set_run_font(run, size=9.2, bold=False, color=DARK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_style_font(style, size, bold=False, color=DARK):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_section_heading(doc, title):
    p = doc.add_paragraph(style="Section Heading")
    p.add_run(title)
    keep_with_next(p)
    return p


def add_label(doc, text):
    p = doc.add_paragraph(style="Block Label")
    p.add_run(text)
    keep_with_next(p)
    return p


def add_verse(doc, number, text):
    p = doc.add_paragraph(style="Verse")
    r = p.add_run(f"{number} ")
    set_run_font(r, size=9.0, bold=True, color=BLUE)
    r = p.add_run(text)
    set_run_font(r, size=9.0)
    return p


def add_bullets(doc, items, point=False):
    style = "Point Bullet" if point else "Background Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(6.5)
    section.bottom_margin = Mm(6.5)
    section.left_margin = Mm(7.5)
    section.right_margin = Mm(7.5)
    section.header_distance = Mm(3.5)
    section.footer_distance = Mm(3.5)

    normal = doc.styles["Normal"]
    set_style_font(normal, 9.0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0.5)
    normal.paragraph_format.line_spacing = Pt(12)

    styles = doc.styles
    for name in ["Section Heading", "Block Label", "Verse", "Background Bullet", "Point Bullet", "Summary"]:
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    s = styles["Section Heading"]
    set_style_font(s, 10.8, bold=True, color=BLUE)
    s.paragraph_format.space_before = Pt(2)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.keep_with_next = True

    s = styles["Block Label"]
    set_style_font(s, 9.0, bold=True, color=BLUE)
    s.paragraph_format.space_before = Pt(1)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.keep_with_next = True

    s = styles["Verse"]
    set_style_font(s, 9.0)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.line_spacing = Pt(12)
    s.paragraph_format.widow_control = True

    for name, size, color in [("Background Bullet", 9.0, MUTED), ("Point Bullet", 9.0, DARK)]:
        s = styles[name]
        set_style_font(s, size, color=color)
        s.base_style = styles["List Bullet"]
        s.paragraph_format.left_indent = Mm(3.5)
        s.paragraph_format.first_line_indent = Mm(-1.8)
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(0.2 if name == "Background Bullet" else 0.5)
        s.paragraph_format.line_spacing = Pt(12)

    s = styles["Summary"]
    set_style_font(s, 9.0, color=DARK)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0.3)
    s.paragraph_format.line_spacing = Pt(12)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("エゼキエル26:1–21　早天メッセージ資料")
    set_run_font(r, size=13.5, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(2)
    r = subtitle.add_run("本文 → 背景・語句 → ポイント → 全体のまとめ → 適用")
    set_run_font(r, size=9.0, color=MUTED)

    sections = [
        {
            "title": "1．エルサレムの滅びを喜ぶツロ（26:1–6）",
            "verses": [
                ("26:1", "第十一年の、その月の一日、私に次のような主のことばがあった。"),
                ("26:2", "「人の子よ。ツロはエルサレムについて、『あはは。国々の民の門は壊され、私に明け渡された。私は豊かになり、エルサレムは廃墟となった』と言った。"),
                ("26:3", "それゆえ──神である主はこう言われる── ツロよ、わたしはおまえを敵とする。海が波をうねらせるように、多くの国々をおまえに向けて攻め上らせる。"),
                ("26:4", "彼らはツロの城壁を荒らし、そのやぐらを壊す。わたしはそのちりを払い去って、そこを裸岩にする。"),
                ("26:5", "ツロは海の中の網干し場となる。わたしが語ったからだ。──神である主のことば── ツロは諸国の餌食となり、"),
                ("26:6", "それに属する沿岸側の町々も剣で滅ぼされる。そのとき彼らは、わたしが主であることを知る。」"),
            ],
            "background": [
                "第十一年はエルサレム陥落の前後。ただし月名が欠けているため、正確な日付は不明",
                "ツロはダビデ、ソロモン時代からイスラエルと建築や交易で結ばれていた（Ⅱサムエル5:11、Ⅰ列王記5:1–12）",
                "「国々の民の門」――エルサレムを交易と交通の重要な入口として見る表現",
                "エルサレムを失った交易が、自分のもとへ流れてくると期待したツロ",
            ],
            "points": [
                "「あはは」――エルサレムの滅びを悲しまず、自分が豊かになる機会として喜ぶツロ",
                "かつて結ばれていた隣人さえ、利益を求める目には競争相手としか見えなくなる",
                "「わたしはおまえを敵とする」――人への冷酷さが、主に敵対する罪として裁かれる",
            ],
        },
        {
            "title": "2．ネブカドネツァルによる攻撃（26:7–14）",
            "verses": [
                ("26:7", "まことに、神である主はこう言われる。「見よ。わたしは、王の王、バビロンの王ネブカドネツァルを、馬、戦車、騎兵、そして大軍勢とともに、北からツロに連れて来る。"),
                ("26:8", "彼はその沿岸側の町々を剣で滅ぼし、おまえに向かって塁を築き、城壁崩しを設け、大盾を立て、"),
                ("26:9", "破城槌でおまえの城壁を突き崩し、やぐらを斧で打ち壊す。"),
                ("26:10", "彼の馬の数があまりにも多いため、その土煙がおまえをおおう。打ち破られた町に入る者のように、彼がおまえの城門に入るとき、騎兵と車両と戦車の響きに、おまえの城壁は揺れ動く。"),
                ("26:11", "彼は、馬のひづめでおまえの大通りをすべて踏みにじり、剣でおまえの民を殺し、おまえの巨大な石柱も地に倒れる。"),
                ("26:12", "彼らはおまえの財宝を略奪し、商品をかすめ奪い、城壁を破壊し、住み心地のよい家を打ち壊し、石や木や土までも、水の中に投げ込む。"),
                ("26:13", "わたしはおまえの騒がしい歌をやめさせる。おまえの竪琴の音も、もう聞かれない。"),
                ("26:14", "わたしはおまえを裸岩とする。おまえは網干し場となり、二度と建て直されない。主であるわたしが語ったからだ。──神である主のことば。」"),
            ],
            "background": [
                "ネブカドネツァル――エルサレムを滅ぼしたバビロン王。ここでは主が連れて来る裁きの道具",
                "「王の王」――ネブカドネツァルの強大な支配を示す称号。主ご自身を指す呼び名ではない",
                "7–11節の単数「彼」から、12節では複数の「彼ら」へ変化。3節の「多くの国々」へ視野が広がる可能性",
                "アレクサンドロスまで含む理解は可能だが本文は彼の名を語らない。29:18も踏まえ、ネブカドネツァル一人が全細部を成就したと単純化しない",
            ],
            "points": [
                "「わたしは連れて来る」――最強の帝国の王も、主の裁きを行うために用いられる",
                "城壁、やぐら、財宝、商品、歌――ツロが誇った力と豊かさと喜びが順に取り去られる",
                "「主であるわたしが語った」――歴史を最後に決めるのは、都市の力でも王の力でもなく主のことば",
            ],
        },
        {
            "title": "3．ツロの崩壊に震える国々（26:15–18）",
            "verses": [
                ("26:15", "神である主はツロにこう言われる。「刺された者がうめき、おまえの中で虐殺が行われるとき、おまえが崩れ落ちるその響きに、島々は揺れ動かないだろうか。"),
                ("26:16", "海の君主たちはみな、その王座から降り、上着を脱ぎ捨て、あや織りの衣服を脱ぐ。彼らは戦慄を身にまとって地面に座り、おまえのことで絶えず身震いし、啞然とする。"),
                ("26:17", "彼らはおまえについて哀歌を唱えて言う。海に住む者よ、おまえはどうして海から消え失せたのか。その町と住民は海で最も強く、ほめそやされた町であったのに。その町の住民すべてに、恐怖がもたらされた。"),
                ("26:18", "今、島々はおまえが崩れ落ちる日に身震いし、海の島々はおまえの退却を見てうろたえる。」"),
            ],
            "background": [
                "「島々」――島だけでなく、地中海沿岸の国々やツロの交易相手まで含む表現",
                "「海の君主たち」――ツロの交易によって繁栄していた沿岸地域の支配者たち",
                "王座から降り、豪華な衣服を脱いで地面に座る姿――栄光と力を失ったしるし",
                "「哀歌」――すでに死んだ者を悼むように、ツロの滅びを歌うもの",
            ],
            "points": [
                "海の君主たちが王座から地面へ――ツロの崩壊によって、周辺諸国の誇りも崩される",
                "ツロ一国の滅びに島々まで震える――繁栄を支え合った国々も無関係ではいられない",
                "海で最も強いとほめられた町への賛美が、その滅びを悲しむ哀歌へ変わる",
            ],
        },
        {
            "title": "4．海の都が地下の国へ下る（26:19–21）",
            "verses": [
                ("26:19", "まことに、神である主はこう言われる。「わたしがおまえを廃墟の町とし、住む者のない町々のようにするとき、大水をおまえの上に湧き上がらせ、豊かな水がおまえをおおうとき、"),
                ("26:20", "わたしはおまえを、穴に下った者たちとともに昔の民のもとに下らせ、穴に下った者たちとともに、昔から廃墟であったような地下の国に住まわせる。わたしが誉れを与える生ける者の地に、おまえが住めないようにするためだ。"),
                ("26:21", "わたしはおまえを恐怖のもととする。おまえはもう存在しなくなり、人がおまえを尋ねても、永久におまえを見つけることはない──神である主のことば。」"),
            ],
            "background": [
                "「大水」――ツロに富を運んだ海が、今度はツロをのみ込む裁きの姿として描かれる",
                "「穴に下る」――単なる都市の敗北を超え、死者の世界へ下る姿",
                "「昔の民」――すでに死んで地下の国にいる過去の人々",
                "「生ける者の地」――主が命と誉れを与える地上の世界",
            ],
            "points": [
                "海によって豊かになったツロが、その海に覆われ、死者のいる地下の国へ下っていく",
                "「生ける者の地」に住めない――命も誉れも、人が自分の力で確保できるものではない",
                "人々に恐れられ、ほめられた名も、主の裁きの前では永久に残ることができない",
            ],
        },
    ]

    for sec in sections:
        add_section_heading(doc, sec["title"])
        add_label(doc, "本文")
        for number, text in sec["verses"]:
            add_verse(doc, number, text)
        add_label(doc, "背景・語句")
        add_bullets(doc, sec["background"], point=False)
        add_label(doc, "ポイント")
        add_bullets(doc, sec["points"], point=True)

    add_section_heading(doc, "全体のまとめ")
    for text in [
        "他者の滅びを自分の利益とし、富と力を神のように頼ったツロ。",
        "主は諸国の歴史を用いてその誇りを倒し、裁きを通してご自身を知らせる。",
        "命、誉れ、繁栄を与え、それを取り去る主だけが、国々の歴史の上に立つ。",
    ]:
        p = doc.add_paragraph(style="Summary")
        p.add_run(text)

    add_label(doc, "中心命題")
    p = doc.add_paragraph(style="Summary")
    shade_paragraph(p, "EEF3F8")
    p.paragraph_format.left_indent = Mm(2)
    p.paragraph_format.right_indent = Mm(2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("人の繁栄を最後に支えるのは、富や力ではなく、すべてを与える主との関係である。")
    set_run_font(r, size=9.2, bold=True, color=BLUE)

    add_section_heading(doc, "適用")
    applications = [
        "「あはは」――誰かの失敗や喪失を、自分の利益や立場の向上として喜んでいないか",
        "かつての隣人――長い関係があっても、利益を中心にすると相手の痛みが見えなくなる",
        "城壁と財宝――自分を守るために築いたものを、主より確かなものとして頼っていないか",
        "歌が止まる――順調な生活に支えられた喜びは、それが失われるとともに消えてしまう",
        "「わたしは敵とする」――敵であった者を、イエス様の十字架によって和解へ招く神（ローマ5:10）",
    ]
    add_bullets(doc, applications, point=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("※適用5はエゼキエル26章の直接預言ではなく、聖書全体からの福音的適用。")
    set_run_font(r, size=9.0, color=MUTED)

    core = doc.core_properties
    core.title = "エゼキエル26:1–21 早天メッセージ資料"
    core.subject = "本文・背景と語句・釈義ポイント・まとめ・適用"
    core.author = ""
    core.keywords = "エゼキエル, 早天, 説教, 釈義"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
