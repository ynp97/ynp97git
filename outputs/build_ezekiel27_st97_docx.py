from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "outputs/20260720_エゼキエル27章1-25節_早天メッセージ.docx"
FONT_JA = "Meiryo UI"
FONT_LATIN = "Kozuka Mincho Pro L"
BLUE = "006699"
DARK = "000000"
LIGHT_BLUE = "EAF3F7"
PALE_GOLD = "FFF6DD"
GRAY = "59636B"


def set_run_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_JA)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(p, before=0, after=3, line=15, keep_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    pf.keep_with_next = keep_next
    pf.widow_control = True


def shade_paragraph(p, fill, border=None):
    ppr = p._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        pbdr = ppr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr")
            ppr.append(pbdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), border)
        pbdr.append(left)


def set_cell_margins(p, left=1.5, right=1.5):
    p.paragraph_format.left_indent = Mm(left)
    p.paragraph_format.right_indent = Mm(right)


doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Mm(35)
sec.bottom_margin = Mm(30)
sec.left_margin = Mm(30)
sec.right_margin = Mm(30)
sec.header_distance = Mm(15)
sec.footer_distance = Mm(17.5)

# Compact reference guide, with named st97 overrides required by the Vault spec.
normal = doc.styles["Normal"]
normal.font.name = FONT_LATIN
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JA)
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = Pt(15)

for style_name, size, color, before, after in [
    ("Title", 16, DARK, 0, 5),
    ("Heading 1", 13, DARK, 9, 4),
    ("Heading 2", 13, BLUE, 8, 3),
    ("Heading 3", 10.5, BLUE, 5, 2),
]:
    s = doc.styles[style_name]
    s.font.name = FONT_LATIN
    s._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    s._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JA)
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.line_spacing = Pt(15)
    s.paragraph_format.keep_with_next = True

if "Scripture" not in doc.styles:
    scripture_style = doc.styles.add_style("Scripture", WD_STYLE_TYPE.PARAGRAPH)
else:
    scripture_style = doc.styles["Scripture"]
scripture_style.font.name = FONT_LATIN
scripture_style.font.size = Pt(10.5)
scripture_style.font.bold = True
scripture_style.paragraph_format.space_before = Pt(3)
scripture_style.paragraph_format.space_after = Pt(3)
scripture_style.paragraph_format.line_spacing = Pt(15)
scripture_style.paragraph_format.keep_together = True


def add_title(text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, 0, 5, 19, True)
    r = p.add_run(text)
    set_run_font(r, 16, True, DARK)
    if subtitle:
        sp = doc.add_paragraph()
        set_spacing(sp, 0, 7, 15, True)
        sr = sp.add_run(subtitle)
        set_run_font(sr, 10.5, False, GRAY)


def h1(text):
    return doc.add_paragraph(text, style="Heading 1")


def h2(text):
    return doc.add_paragraph(text, style="Heading 2")


def h3(text):
    p = doc.add_paragraph()
    set_spacing(p, 5, 5, 15, True)
    r = p.add_run(text)
    set_run_font(r, 10.5, True, BLUE)
    if text == "段落のポイント":
        r.add_break()
    return p


def add_bullet(text, level=0, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Mm(5.5 + level * 4)
    p.paragraph_format.first_line_indent = Mm(-3.2)
    set_spacing(p, 1.5, after, 15)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.first_line_indent = Mm(-3.5)
    set_spacing(p, 0, 2, 15)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_body(text, bold=False, fill=None, color=None, after=3):
    p = doc.add_paragraph()
    set_spacing(p, 0, after, 15)
    r = p.add_run(text)
    set_run_font(r, 10.5, bold, color)
    if fill:
        shade_paragraph(p, fill, BLUE)
        set_cell_margins(p)
    return p


def add_verse(num, text, notes):
    p = doc.add_paragraph(style="Scripture")
    set_spacing(p, 3, 3, 15, True)
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Mm(6.35)
    nr = p.add_run(f"エゼキエル{num}　")
    set_run_font(nr, 10.5, True, BLUE)
    tr = p.add_run(text)
    set_run_font(tr, 10.5, True, "000000")
    label = doc.add_paragraph()
    set_spacing(label, 0, 1, 15, True)
    lr = label.add_run("背景・語句")
    set_run_font(lr, 10.5, True, BLUE)
    for note in notes:
        add_bullet(note)


add_title("2026年7月20日　エゼキエル27:1–25\n「満ちあふれ、重くなった船」")

h1("全体の流れと結論")
for item in [
    "主はエゼキエルに、繁栄するツロのための「哀歌」を命じる。",
    "ツロは、自分を「美の極み」と誇る豪華な商船として描かれる。",
    "世界各地の材料、人材、軍事力、商品がツロに集まり、その繁栄は頂点に達する。",
    "しかし、これは成功を祝う歌ではなく、滅びを先取りした哀歌である。続く26節以降で、この船は海に沈む。",
]:
    add_numbered(item)
add_body("結論：神から受けた豊かさを自分の完全さとして誇るなら、その豊かさ自体が、神に頼らない生き方の重荷となる。", True, PALE_GOLD, DARK, 6)
add_body("27章はツロを巨大な商船として描く。9節付近から船と都市の表現が重なり、26節以降では船の沈没によってツロの滅亡が語られる。", False, None, GRAY, 5)


sections = [
    (
        "第一段落　哀歌として語られるツロの美しさ（1〜4節）",
        "ツロは海上交易によって栄え、自らを「美の極み」と評価していた。しかし主は、その繁栄の最中にツロのための哀歌を命じる。",
        [
            ("27:1", "次のような主のことばが私にあった。", [
                "26〜28章は、ツロに対する一連の預言。",
                "26章では、ツロがエルサレムの滅亡を自分の利益として喜んだことが問題にされている。",
                "27章はツロの繁栄と崩壊、28章はツロの指導者の高慢を扱う。",
            ]),
            ("27:2", "「人の子よ、ツロについて哀歌を唱えよ。", [
                "「哀歌」はヘブル語でキーナー。通常は死者や滅びた町を悼む歌。",
                "この時点のツロはまだ大いに栄えている。それなのに哀歌が命じられることが、この章の緊張を生む。",
                "神の判定では、外から繁栄して見えるツロの滅亡はすでに確実なものとして語られている。",
            ]),
            ("27:3", "あなたはツロに言え。\n『海の出入り口に住み、\n多くの島々の民と取り引きをする者よ、\n神である主はこう言われる。\nツロよ、おまえは、「私は美の極みだ」と言った。", [
                "「海の出入り口」は、地中海交易の重要な拠点であったツロの位置と役割を表す。複数形の原文は、ツロの二つの港を指す可能性もある。",
                "「島々」は島だけでなく、地中海沿岸の遠い国々を広く含む表現。",
                "「美の極み」は、欠けたところのない完全な美という意味。28章でもツロの支配者の高慢と結びつく。",
                "問題は、美しさや商業そのものより、ツロがそれを自分の完全さとして語ったことにある。",
            ]),
            ("27:4", "おまえの領土は海の真ん中にあり、\nおまえを築いた者は、\nおまえを最高に美しく仕上げた。", [
                "ここからツロは、一隻の豪華な船として描かれる。",
                "「海の真ん中」は、海上の島を中心としていたツロの地理と、その国際的な活動を重ねた表現。",
                "ツロの美しさは単なる思い込みではない。実際に多くの人々の技術によって、美しく築かれていた。",
                "しかしツロは、多くの土地と人から受け取った美しさを、自分自身の完全さとして誇った。",
            ]),
        ],
        [
            "主はツロの繁栄を否定していない。実際の美しさと成功を認めたうえで、その結末を哀歌として語る。",
            "ツロは多くの人々に支えられていたのに、「私は美の極みだ」と自分を繁栄の中心に置いた。",
            "人間が成功と神の承認を同じものと考えるとき、現在の繁栄によって迫っている危険が見えなくなる。",
        ],
    ),
    (
        "第二段落　世界中の材料と人材で造られた船（5〜11節）",
        "ツロという船には、各地から最高の木材、布、技術者、船員、兵士が集められた。その美しさと強さは、国際的な協力によって造られていた。",
        [
            ("27:5", "彼らはセニルのもみの木で\nおまえのすべての船板を作り、\nレバノンの杉を使って、おまえの帆柱を作った。", [
                "セニルはヘルモン山周辺を指す名称。",
                "レバノン杉は、大きさと耐久性で知られた高級建材。",
                "船体から帆柱まで、各地の優れた木材が使われたという表現によって、ツロの豊かさが示される。",
            ]),
            ("27:6", "バシャンの樫の木でおまえの櫂を作り、\nキティムの島々の檜に象牙をはめ込んで、\nおまえの甲板を作った。", [
                "バシャンは良質な樫の木で知られたヨルダン川東方の地域。",
                "キティムは、おもにキプロス周辺を指すと考えられる。",
                "この節のヘブル語には本文上の難しさがあり、甲板、船室、船体など訳が分かれる。象牙をはめ込んだ豪華な造りという中心的な意味は明確である。",
            ]),
            ("27:7", "エジプトのあや織りの亜麻布が、\nおまえの帆であり、\nおまえの旗じるしであった。\nエリシャの島々からの青色と紫色の布が、\nおまえの覆いであった。", [
                "エジプトの上質な亜麻布は、本来実用的な帆を、遠くからも見える美しい旗じるしにした。",
                "青色と紫色の布は、高価な染料を使った富と高い身分の象徴。",
                "エリシャの正確な位置は確定できないが、キプロスまたは地中海沿岸の地域と考えられている。",
                "ツロは荷物を運ぶだけの船ではなく、自分の豊かさを見せる船として描かれる。",
            ]),
            ("27:8", "シドンとアルワデの住民が、\nおまえの漕ぎ手であった。\nツロよ、おまえのうちの熟練者が、\nおまえの船員であった。", [
                "シドンとアルワデは、フェニキア沿岸の海上活動で知られた都市。",
                "周辺諸国の人々が漕ぎ手となり、ツロ自身の熟練者が船を指揮した。",
                "ツロの成功には、資材だけでなく、各地の労働力と専門技術が必要だった。",
            ]),
            ("27:9", "ゲバルの長老と、その熟練者が、\nおまえのうちにあって破損を修理し、\n海のすべての船とその水夫たちが、\nおまえのうちにあって\nおまえに商品を持ち込んだ。", [
                "ゲバルは後のビブロス。造船と海上交易で知られたフェニキアの都市。",
                "「破損を修理し」は、船板の継ぎ目をふさぎ、船を航行可能に保つ熟練作業。",
                "後半では「ツロという船」から「船々が入港する都市ツロ」へと表現が移り始める。",
                "ツロは、自分だけで富を生み出すのではなく、諸国の商品が集まる市場だった。",
            ]),
            ("27:10", "ペルシア、ルデ、プテの人々は、\nおまえの軍隊の戦士であり、\nおまえに盾とかぶとを掛け、\n彼らはおまえに輝きを添えた。", [
                "ペルシアは東方、ルデは小アジア方面、プテは北アフリカ方面と考えられる。",
                "ツロは外国人兵士を用いて、自国の安全を確保していた。",
                "壁に掛けられた盾とかぶとは、防衛力であると同時に都市を飾る威容でもあった。",
            ]),
            ("27:11", "アルワデとヘレクの人々はおまえの周りの城壁の上に、またガマデ人はおまえのやぐらの中にいて、周りの城壁に丸い小盾を掛け、おまえを最高に美しく仕上げた。", [
                "ヘレクとガマデ人の正確な特定には不確かさが残る。",
                "城壁、やぐら、外国人兵士、小盾によって、ツロは安全で堂々とした都市に見えた。",
                "3〜4節の「美の極み」「最高に美しく仕上げた」という表現が、ここで繰り返される。",
                "その美しさには、軍事力と安全保障も含まれていた。",
            ]),
        ],
        [
            "ツロの「美」は、セニル、レバノン、バシャン、キティム、エジプトなど、諸地域から受け取ったものの結集だった。",
            "外国の材料、職人、船員、兵士がいなければ、ツロという船は完成しなかった。",
            "受け取ったものを数えれば感謝に向かう。しかし、それを自分の完全さの証明にすると高慢に向かう。",
        ],
    ),
    (
        "第三段落　世界を結びつけた巨大な商業網（12〜19節）",
        "ツロには、西方、小アジア、シリア、イスラエルなどから、鉱物、人間、家畜、食料、宝石、香料が集まった。ツロは地中海世界の大きな市場となっていた。",
        [
            ("27:12", "タルシシュは、おまえがあらゆる財宝に富んでいたので、おまえと商いをし、銀、鉄、すず、鉛をおまえの商品と交換した。", [
                "タルシシュの正確な位置には議論があるが、遠い西方の交易地として描かれている。",
                "銀、鉄、すず、鉛は、ツロが鉱物資源を広範囲から集めていたことを示す。",
                "「あらゆる財宝に富む」という評価が、ツロの経済的な大きさを表す。",
            ]),
            ("27:13", "ヤワン、トバル、メシェクはおまえと取り引きをし、人間と青銅の器具をおまえの商品と交換した。", [
                "ヤワンはギリシア・イオニア方面、トバルとメシェクは小アジア方面の民と考えられる。",
                "「人間」は奴隷として売買された人々を指す。",
                "この節は、ツロの繁栄が華やかな商品だけでなく、人間を商品にする取引も含んでいたことを示す。",
                "本文はここで詳しい道徳的評価を加えないが、交易の暗い面を隠してはいない。",
            ]),
            ("27:14", "ベテ・トガルマは馬、軍馬、らばをおまえの商品と交換した。", [
                "ベテ・トガルマは小アジア東部、アルメニア方面と考えられる。",
                "馬、軍馬、らばは、移動、輸送、軍事に必要な高価な資源だった。",
            ]),
            ("27:15", "デダン人はおまえと取り引きをし、多くの島々はおまえの支配する市場であり、彼らは象牙と黒檀をおまえに貢ぎとして持って来た。", [
                "ここにある「デダン」の場所については、写本や訳によって理解が分かれる。",
                "象牙と黒檀は、アフリカや南方との遠距離交易を示す高級品。",
                "「貢ぎ」は、単なる売買だけでなく、ツロが市場を支配する強い立場にあったことを示す。",
            ]),
            ("27:16", "アラムは、おまえの事業が多岐にわたったので、おまえと商いをし、トルコ石、紫色の布、あや織物、白亜麻布、珊瑚、紅玉をおまえの商品と交換した。", [
                "アラムはシリア方面。",
                "ここに並ぶのは宝石、染色布、上質な織物などの高級品。",
                "「事業が多岐にわたった」は、ツロが扱う商品の豊富さを表す。",
                "いくつかの商品名は古代語として特定が難しく、訳語には幅がある。",
            ]),
            ("27:17", "ユダとイスラエルの地もおまえと取り引きをし、ミニテの小麦、きび、蜜、香油、乳香をおまえの商品と交換した。", [
                "ユダとイスラエルも、ツロの交易網に組み込まれていた。",
                "穀物、蜜、香油、乳香は、土地の産物や生活用品。",
                "「きび」に当たる原語は意味が不確かで、菓子、香料、特定の食料品などの訳もある。",
                "エゼキエルの民も、ツロと無関係な外部の観察者ではなかった。",
            ]),
            ("27:18", "おまえの事業が多岐にわたり、あらゆる財宝に富んでいたので、ダマスコも、ヘルボンのぶどう酒とツァハルの羊毛でおまえと商いをした。", [
                "ダマスコは内陸交易の重要都市。",
                "ヘルボンは上質なぶどう酒で知られた地域。",
                "ツァハルの位置や語の意味は確定していないが、本文では羊毛の産地として挙げられる。",
                "12節の「あらゆる財宝に富む」という表現が繰り返され、繁栄の大きさが強調される。",
            ]),
            ("27:19", "ダンとヤワンもおまえの商品と交換した。その商品の中にはウザルからの銑鉄、桂枝、菖蒲があった。", [
                "この節の地名と文章には本文上の難しさがある。「ダン」ではなく「ベダン」、「ウザルから来たヤワン」などと理解する訳もある。",
                "鉄と香料が並び、実用品からぜいたく品まで、ツロの取引範囲が広かったことを示す。",
                "地名の特定は難しくても、遠方から多種類の商品が集まったという本文の役割は明確である。",
            ]),
        ],
        [
            "地名と商品が繰り返される長い一覧は、ツロの影響が非常に広かったことを読者に実感させる。",
            "交易は多くの地域を結び、生活を豊かにした一方、人間まで商品として扱う構造も含んでいた。",
            "この箇所だけから「商売や富そのものが罪だ」と結論づけるべきではない。問題の中心は、富を自分の完全さと安全の根拠にしたことである。",
        ],
    ),
    (
        "第四段落　繁栄の頂点に達したツロ（20〜25節）",
        "アラビアや東方の国々もツロの市場に加わり、織物、家畜、香料、宝石、金が集まった。ツロという船は、商品を満載して大海のただ中で繁栄の頂点に達する。",
        [
            ("27:20", "デダンは鞍に敷く織り布でおまえと取り引きをした。", [
                "ここでのデダンは、アラビア方面の交易民を指すと考えられる。",
                "鞍用の布は、砂漠地帯の隊商と移動生活に関係する商品。",
            ]),
            ("27:21", "アラビア人、ケダルの君主たちもみな、おまえの御用商人であり、子羊、雄羊、やぎで商いをした。", [
                "ケダルは北アラビアの遊牧民として知られる。",
                "家畜による取引は、ツロの市場が都市の高級品だけでなく、牧畜地域の産物とも結ばれていたことを示す。",
                "「御用商人」は、継続的にツロへ商品を供給する商人を表す。",
            ]),
            ("27:22", "シェバとラアマの商人たちはおまえと取り引きをし、あらゆる上等の香料、宝石、また金をおまえの商品と交換した。", [
                "シェバとラアマは南アラビア方面と結びつけられる。",
                "香料、宝石、金は、遠距離交易を代表する高価な商品。",
                "「あらゆる上等の」という表現が、ツロには最良の品が集まったことを強調する。",
            ]),
            ("27:23", "ハラン、カンネ、エデン、それにシェバの商人たち、アッシュルとキルマデはおまえと取り引きをした。", [
                "ハランとアッシュルはメソポタミア方面。カンネ、エデン、キルマデの正確な位置には不確かさがある。",
                "地中海沿岸だけでなく、内陸の交易路もツロにつながっていた。",
                "確認できる地名と不明な地名が混在するため、現代の国名へ一律に置き換えることは避ける。",
            ]),
            ("27:24", "おまえの市場で、彼らは豪華な衣服、青色の衣、あや織物、固く撚った綱でしっかり留められた多彩な敷き物をもって、おまえと取り引きをした。", [
                "豪華な衣服、染色された布、多彩な敷き物は、高度な加工技術と富裕層の需要を示す。",
                "ツロは原材料を集めるだけでなく、各地の完成品が並ぶ国際市場だった。",
            ]),
            ("27:25", "タルシシュ船がおまえの商品を運んだ。\nおまえは大海のただ中で\n満ちあふれて、大いに栄えた。", [
                "「タルシシュ船」は、タルシシュに所属する船だけでなく、遠洋航海に用いられる大型商船を指す場合がある。",
                "「満ちあふれて」は、ツロという船が商品を満載している姿。",
                "ここで1〜25節の繁栄は頂点に達する。",
                "しかし冒頭ですでに「哀歌」と宣言されている。続く26節では、この満載の船が大波の中で砕かれる。",
                "したがって25節は安心の結論ではなく、崩壊直前の頂点である。",
            ]),
        ],
        [
            "ツロは西から東、海洋から内陸、農耕地から遊牧地までを結ぶ巨大な市場だった。",
            "「満ちあふれて、大いに栄えた」は事実である。しかし、豊かさの量は将来の安全を保証しなかった。",
            "船が大きく、積荷が多いほど、それを支える船体と航路が失われたときの損失も大きくなる。",
        ],
    ),
]

for heading, summary, verses, points in sections:
    h1(heading)
    h3("段落の簡単なまとめ")
    add_body(summary, False, PALE_GOLD, DARK, 4)
    for num, text, notes in verses:
        add_verse(num, text, notes)
    h3("段落のポイント")
    for point in points:
        add_bullet(point, after=2)

h1("全体のまとめ")
for item in [
    "ツロは、世界各地の材料、人材、軍事力、商品によって、実際に美しく豊かな都市となった。",
    "しかしツロは、受け取った豊かさを「私は美の極みだ」という自己評価に変えた。",
    "本文は商業、技術、美、富そのものを罪としてはいない。同時に、奴隷取引を含む繁栄の暗い面も記録している。",
    "この章は最初から「哀歌」である。繁栄の頂点を詳しく描くことによって、目に見える成功が永続する保証にはならないことを示す。",
    "ツロという満載の船は、26節以降で沈む。神を離れた自己完結と安全の誇りは、世界規模の繁栄によっても支えきれない。",
]:
    add_bullet(item)

h1("五段階の適用")
applications = [
    ("1．生活と行動", "今日、自分に与えられている能力、働き、人、財産の中から一つを選び、「自分が築いたもの」と言う前に、誰を通して神から与えられたかを書き出して感謝する。"),
    ("2．心の診断と福音の構造", "ツロは、外から集められた美しさを「私は美の極みだ」と言い換えた。私も、奉仕、仕事、知識、経験を、自分の価値を証明する荷物にしていないか。福音は、積み荷の多さによって私の価値が決まるとは言わない。神の恵みが先に与えられ、働きはその恵みへの応答となる。"),
    ("3．キリスト論・終末論――聖書全体からの適用", "黙示録18章はエゼキエル27章の言葉と構造を用いて、世界の商人を富ませた「大バビロン」の滅びを描く。その後に現れるのは、自分で美を築いた都ではなく、「神の栄光」を持つ新しいエルサレムである。イエス様は、人を商品にし、自分を神のように高くする世界を終わらせ、神の栄光を受けて輝く都を完成される。錨：エゼキエル27章／黙示録18:11–19／21:10–11――非常に強。思弁度：低。"),
    ("4．三位一体論・観想――聖書全体からの神学的観想", "父なる神は、すべての良いものの与え主である。御子は神の栄光を持ちながら、それを自分の利益のために握りしめず、私たちのためにご自分を低くされた。聖霊は、受け取った賜物を自己賛美ではなく、他者に仕えるために用いるよう導かれる。宗教的な奉仕も積み荷になり得る。奉仕の実績によって自分の重要性を証明しているなら、私は神に仕えながら、ツロと同じ自己完結へ向かっている。錨：ヤコブ1:17／ピリピ2:6–8／Ⅰコリント12:4–7――中。思弁度：中。"),
    ("5．説教者自身の生", "私の船に今、感謝して受け取っている荷物ではなく、自分の価値を証明するために積み続けている荷物は何か。"),
]
for label, body in applications:
    p = doc.add_paragraph()
    set_spacing(p, 4, 1, 15, True)
    r = p.add_run(label)
    set_run_font(r, 10.5, True, BLUE)
    add_body(body, False, None, DARK, 3)

p = doc.add_paragraph()
set_spacing(p, 7, 0, 15)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("ハルシネーションチェック：再確認済み")
set_run_font(r, 9, False, GRAY)

# Centered page-number footer.
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(fp, 0, 0, 12)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Ensure all paragraphs use Meiryo and headings never strand at a page bottom.
for p in doc.paragraphs:
    for run in p.runs:
        size = run.font.size.pt if run.font.size else 10.5
        bold = bool(run.bold)
        italic = bool(run.italic)
        color = None
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
        set_run_font(run, size, bold, color, italic)
    if p.style and p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True

doc.core_properties.title = "2026年7月20日 エゼキエル27:1–25 満ちあふれ、重くなった船"
doc.core_properties.subject = "早天メッセージ準備ガイド"
doc.core_properties.author = ""
doc.save(OUT)
print(OUT)
