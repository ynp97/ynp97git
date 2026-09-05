# -*- coding: utf-8 -*-
"""st97（早天）のペイロードJSON → Word（.docx）

  python3 scripts/st97_docx_build.py .agents/skills/st97/tmp/20260904_psalm119_49-64.json

出力先は既定で Vault の output/ 直下（第2引数で変更できる）。
書式は 📄 Word出力仕様（AI共通）.md ／ 🌅 早天Word出力仕様（AI共通）.md に従い、
実装は scripts/docx_style.py（set97の釈義Word化と同じ土台）を共有する。

内容の順序は 🌅 早天Word出力仕様 のとおり、印刷用HTML（render_st97.py）と同じ:
  表題 → 全体の流れと結論 → 段落の簡単なまとめ → 各節の本文 → 直下の背景・語句
      → 段落のポイント → 全体のまとめ → 五段階の適用

ペイロードの ruby は Word のルビ（w:ruby）として入れる。
必要なもの: python3 と python-docx（無ければ pip3 install python-docx）
"""
import json, os, sys, re
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx.shared import Pt, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx_style as S

# ---------- ルビ ----------
def _ruby_run(p, base, kana, bold=False, color=None, size_pt=None):
    """<w:r><w:ruby>…</w:ruby></w:r> を段落へ足す。"""
    r = OxmlElement('w:r')
    ruby = OxmlElement('w:ruby')
    pr = OxmlElement('w:rubyPr')
    base_hps = str(int(round((size_pt or 10.5) * 2)))          # 半ポイント
    rt_hps   = str(max(8, int(round((size_pt or 10.5)))))        # ルビは約半分
    raise_hp = str(int(round((size_pt or 10.5) * 1.75)))
    for tag, val in (('w:rubyAlign', 'distributeSpace'), ('w:hps', rt_hps),
                     ('w:hpsRaise', raise_hp), ('w:hpsBaseText', base_hps), ('w:lid', 'ja-JP')):
        e = OxmlElement(tag); e.set(qn('w:val'), val); pr.append(e)
    ruby.append(pr)

    def _seg(tag, text, sz):
        seg = OxmlElement(tag)
        rr  = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        s = OxmlElement('w:sz'); s.set(qn('w:val'), sz); rpr.append(s)
        if bold: rpr.append(OxmlElement('w:b'))
        if color is not None:
            c = OxmlElement('w:color'); c.set(qn('w:val'), str(color)); rpr.append(c)
        rr.append(rpr)
        t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
        rr.append(t); seg.append(rr)
        return seg

    ruby.append(_seg('w:rt', kana, rt_hps))
    ruby.append(_seg('w:rubyBase', base, base_hps))
    r.append(ruby)
    p._p.append(r)


def add_ruby_text(p, text, ruby_map, bold=False, color=None, size=None):
    """ruby_map の見出し語をルビ付きに、それ以外は普通のランで流し込む。"""
    if not ruby_map:
        S.setfont(p.add_run(text), size, bold, color); return p
    keys = sorted(ruby_map, key=len, reverse=True)
    pat = re.compile('|'.join(re.escape(k) for k in keys))
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            S.setfont(p.add_run(text[pos:m.start()]), size, bold, color)
        _ruby_run(p, m.group(0), ruby_map[m.group(0)], bold=bold, color=color, size_pt=size)
        pos = m.end()
    if pos < len(text):
        S.setfont(p.add_run(text[pos:]), size, bold, color)
    if not p.runs and pos == 0:
        S.setfont(p.add_run(text), size, bold, color)
    return p


def lines(v):
    return str(v).splitlines() or ['']


def build(payload_path, out_path=None, template=None):
    data = json.load(open(payload_path, encoding='utf-8'))
    rb   = data.get('ruby') or {}
    doc  = S.open_doc(template)

    # ---- 表題（日付＋箇所＋題）----
    for i, ln in enumerate(lines(data.get('title') or data.get('document_title', ''))):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = S.SPACE if i == 0 else Emu(139700)
        S.setfont(p.add_run(ln), 16, True)

    # ---- 全体の流れと結論 ----
    S.para(doc, 'Heading 1', space=False).add_run('全体の流れと結論')
    for i, item in enumerate(data.get('flow', []), 1):
        p = S.para(doc, indent=S.INDENT)
        add_ruby_text(p, f'{i}. {item}', rb)
    p = S.para(doc)
    S.setfont(p.add_run('結論：'), bold=True)
    add_ruby_text(p, str(data.get('conclusion', '')), rb, bold=True)
    if data.get('overview'):
        p = S.para(doc)
        add_ruby_text(p, str(data['overview']), rb, color=S.GREY)

    # ---- 段落ごと ----
    for sec in data.get('sections', []):
        p = S.para(doc, 'Heading 1', space=False)
        add_ruby_text(p, sec['heading'], rb)

        S.para(doc, 'Heading 2', space=False).add_run('段落の簡単なまとめ')
        for ln in lines(sec.get('summary', '')):
            add_ruby_text(S.para(doc), ln, rb)

        for verse in sec.get('verses', []):
            # 当日の中心テキスト：ラベル青＋本文は黒の太字・一段字下げ
            p = S.para(doc, indent=S.INDENT)
            S.setfont(p.add_run(verse['label']), S.VERSE_PT, bold=True, color=S.BLUE)
            S.setfont(p.add_run('　'),            S.VERSE_PT, bold=True, color=S.BLUE)
            add_ruby_text(p, ' '.join(lines(verse['text'])), rb, bold=True, size=S.VERSE_PT)
            if verse.get('notes'):
                S.para(doc, 'Heading 2', space=False).add_run('背景・語句')
                for n in verse['notes']:
                    add_ruby_text(S.para(doc, indent=S.INDENT), '・' + n, rb)

        if sec.get('points'):
            S.para(doc, 'Heading 2', space=False).add_run('段落のポイント')
            for pt in sec['points']:
                add_ruby_text(S.para(doc, indent=S.INDENT), '・' + pt, rb)

    # ---- 全体のまとめ ----
    if data.get('overall'):
        S.para(doc, 'Heading 1', space=False).add_run('全体のまとめ')
        for item in data['overall']:
            add_ruby_text(S.para(doc, indent=S.INDENT), '・' + item, rb)

    # ---- 五段階の適用 ----
    if data.get('applications'):
        S.para(doc, 'Heading 1', space=False).add_run('五段階の適用')
        for app in data['applications']:
            p = S.para(doc, 'Heading 2', space=False)
            add_ruby_text(p, app['label'], rb)
            for ln in lines(app['text']):
                add_ruby_text(S.para(doc, indent=S.INDENT), ln, rb)

    if out_path is None:
        base = os.path.splitext(os.path.basename(payload_path))[0]
        vault = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        outdir = os.path.join(vault, 'output')
        os.makedirs(outdir, exist_ok=True)
        out_path = os.path.join(outdir, base + '_早天メッセージ.docx')
    doc.save(out_path)
    return out_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        raise SystemExit(__doc__)
    out = build(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None,
                sys.argv[3] if len(sys.argv) > 3 else None)
    print('保存しました:', out)
