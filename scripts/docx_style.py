# -*- coding: utf-8 -*-
"""Vault共通のWord書式ヘルパー（📄 Word出力仕様（AI共通）.md の実装）

用紙・余白・フォント・見出し・フッターのページ番号は
scripts/vault_docx_template.docx（本人の実原稿から作った空テンプレート）が持っている。
このモジュールは、その上に置く「色分け・字下げ・表組み」だけを担当する。

  釈義する本文   … 節ラベル 青 #006699（太字）＋ 本文テキスト 黒の太字、一段字下げ
  引用・参照聖句 … ラベル・本文とも オレンジ #ED7D31、一段字下げ
  ESV           … グレー #595959、一段字下げ
  表            … 罫線 #A6B4C2 sz4 ／ 幅8504dxa ／ 固定レイアウト
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BLUE   = RGBColor(0x00, 0x66, 0x99)   # 当日の中心テキストの節ラベル
GREY   = RGBColor(0x59, 0x59, 0x59)   # ESV・副題
ORANGE = RGBColor(0xED, 0x7D, 0x31)   # 引用・参照聖句
INDENT = Emu(228600)                  # 聖書の一段字下げ（約0.25インチ）
SPACE  = Emu(63500)                   # 段落後の間隔
TBL_W  = 8504                         # 表の幅（dxa）
VERSE_PT = 10.5                       # 聖書本文の文字サイズ（📄 Word出力仕様どおり本文と同じ）

HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE = os.path.join(HERE, "vault_docx_template.docx")

INLINE = re.compile(r'\*\*(.+?)\*\*')


def open_doc(template=None):
    """空テンプレートを開く。本文が残っていれば消してから返す。"""
    tpl = template or DEFAULT_TEMPLATE
    doc = Document(tpl)
    body = doc.element.body
    for ch in list(body):
        if ch.tag == qn('w:sectPr'):
            continue
        body.remove(ch)
    return doc


def setfont(run, size=None, bold=None, color=None):
    if size  is not None: run.font.size = Pt(size)
    if bold  is not None: run.bold = bold
    if color is not None: run.font.color.rgb = color


def strip_md(t):
    t = INLINE.sub(r'\1', t)
    t = re.sub(r'`([^`]+)`', r'\1', t)
    t = re.sub(r'\[\[(.+?)\]\]', r'\1', t)
    return t


def add_rich(p, text, base_bold=False, color=None, size=None):
    """**強調** を太字ランに分解して段落へ流し込む。"""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            setfont(p.add_run(text[pos:m.start()]), size, base_bold, color)
        setfont(p.add_run(m.group(1)), size, True, color)
        pos = m.end()
    if pos < len(text):
        setfont(p.add_run(text[pos:]), size, base_bold, color)
    if not p.runs:
        setfont(p.add_run(""), size, base_bold, color)
    return p


def para(doc, style=None, indent=None, space=True):
    p = doc.add_paragraph()
    if style: p.style = doc.styles[style]
    if indent is not None: p.paragraph_format.left_indent = indent
    if space: p.paragraph_format.space_after = SPACE
    return p


def title_block(doc, title, subtitle=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = SPACE
    setfont(p.add_run(title), 16, True)
    if subtitle:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Emu(139700)
        add_rich(p, subtitle, color=GREY)


def main_verse(doc, label, text):
    """当日の中心テキスト：ラベル青・本文は黒の太字・一段字下げ・本文より一段大きく。"""
    p = para(doc, indent=INDENT)
    setfont(p.add_run(label), VERSE_PT, bold=True, color=BLUE)
    setfont(p.add_run('　'),  VERSE_PT, bold=True, color=BLUE)
    setfont(p.add_run(strip_md(text)), VERSE_PT, bold=True)
    return p


def quoted_verse(doc, label, text):
    """他箇所からの引用・参照聖句：全体オレンジ・一段字下げ・本文より一段大きく。"""
    p = para(doc, indent=INDENT)
    setfont(p.add_run(f'{label}　{strip_md(text)}'), VERSE_PT, color=ORANGE)
    return p


def esv_verse(doc, label, text):
    p = para(doc, indent=INDENT)
    setfont(p.add_run(f'{label}　{text}'), VERSE_PT, color=GREY)
    return p


def style_table(t, ncol):
    """本人の実原稿の表組みに合わせる。列幅のリストを返す。"""
    tblPr = t._tbl.tblPr
    for tag in ('w:tblW', 'w:tblBorders', 'w:tblLayout'):
        for e in tblPr.findall(qn(tag)):
            tblPr.remove(e)
    w = OxmlElement('w:tblW'); w.set(qn('w:w'), str(TBL_W)); w.set(qn('w:type'), 'dxa'); tblPr.append(w)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0');    e.set(qn('w:color'), 'A6B4C2')
        b.append(e)
    tblPr.append(b)
    if ncol == 1:
        widths = [TBL_W]
    else:
        first = int(TBL_W * (0.14 if ncol >= 4 else (0.22 if ncol == 3 else 0.28)))
        rest  = (TBL_W - first) // (ncol - 1)
        widths = [first] + [rest] * (ncol - 1)
        widths[-1] += TBL_W - sum(widths)
    grid = t._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for gc, wd in zip(grid.findall(qn('w:gridCol')), widths):
            gc.set(qn('w:w'), str(wd))
    return widths


def add_table(doc, rows, header_bold=True):
    if not rows: return None
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    widths = style_table(t, ncol)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, cell in enumerate(cells):
            cell.width = Emu(int(widths[ci] * 635))
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Emu(0)
            add_rich(cp, row[ci] if ci < len(row) else '', base_bold=(header_bold and ri == 0))
    doc.add_paragraph()
    return t
