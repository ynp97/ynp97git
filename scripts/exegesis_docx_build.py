# -*- coding: utf-8 -*-
import re, sys, copy
import os as _os
sys.path.insert(0, _os.path.dirname(_os.path.abspath(__file__)))
import docx_style as S
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.oxml.ns import qn

# ---------------------------------------------------------------
# 使い方（引数は「釈義の .md」1つだけでよい）
#
#   python3 scripts/exegesis_docx_build.py "釈義/マルコ3章20-35節 釈義.md"
#
# 出力先・テンプレート・表題・副題は自動で決まる。
#   出力先     … 同じフォルダに同じ名前の .docx
#   テンプレート … 釈義/ にある既存の 釈義.docx のうち一番新しいもの
#                 （余白・フォント・見出しスタイルをそこから受け継ぐ）
#   表題       … 「YYYY年M月D日　マルコの福音書3:20–35　釈義」（実行日＋mdの見出し）
#   副題       … md の3行目の（…）行
#
# 上書きしたいときだけ、続けて  <テンプレdocx> <出力.docx> "<表題>" "<副題>"  を渡す。
# 必要なもの: python3 と python-docx（無ければ  pip3 install python-docx  ）
# ---------------------------------------------------------------
import os, glob, datetime, re as _re

MD = sys.argv[1]

def _auto_template(md_path):
    if len(sys.argv) > 2: return sys.argv[2]
    if os.path.exists(S.DEFAULT_TEMPLATE): return S.DEFAULT_TEMPLATE
    d = os.path.dirname(os.path.abspath(md_path))
    cands = [f for f in glob.glob(os.path.join(d, "*釈義*.docx"))
             if os.path.abspath(f) != os.path.abspath(_auto_out(md_path))]
    if not cands:
        raise SystemExit("テンプレートにする既存の釈義.docxが見つかりません。第2引数で指定してください。")
    return max(cands, key=os.path.getmtime)

def _auto_out(md_path):
    if len(sys.argv) > 3: return sys.argv[3]
    return os.path.splitext(md_path)[0] + ".docx"

def _auto_title_subtitle(md_path):
    head = open(md_path, encoding="utf-8").read().split("\n")[:6]
    h1 = next((l[2:].strip() for l in head if l.startswith("# ")), "釈義")
    sub = next((l.strip() for l in head if l.strip().startswith("（")), "")
    t = datetime.date.today()
    title = f"{t.year}年{t.month}月{t.day}日　{h1.replace(' 釈義','')}　釈義"
    if len(sys.argv) > 4: title = sys.argv[4]
    if len(sys.argv) > 5: sub   = sys.argv[5]
    return title, sub

OUT = _auto_out(MD)
TPL = _auto_template(MD)
TITLE, SUBTITLE = _auto_title_subtitle(MD)
print("入力:", MD)
print("テンプレート:", TPL)
print("出力:", OUT)
print("表題:", TITLE)

BLUE, GREY, ORANGE = S.BLUE, S.GREY, S.ORANGE
INDENT, SPACE = S.INDENT, S.SPACE
setfont, add_rich, strip_md, style_table = S.setfont, S.add_rich, S.strip_md, S.style_table

doc = Document(TPL)
# 本文を空にする（styles / sectPr は残す）
body = doc.element.body
for child in list(body):
    if child.tag == qn('w:sectPr'):
        continue
    body.remove(child)



def newp(style=None):
    p = doc.add_paragraph()
    if style: p.style = doc.styles[style]
    return p



# ---------- 表題 ----------
p = newp(); p.paragraph_format.space_after = SPACE
r = p.add_run(TITLE); setfont(r, 16, True)
p = newp(); p.paragraph_format.space_after = Emu(139700)
add_rich(p, SUBTITLE, color=GREY)

lines = open(MD, encoding='utf-8').read().split('\n')

quote_mode = None   # 'esv' / 'sky' / ('ref', ラベル書名)
i = 0
def flush_table(rows):
    if not rows: return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    style_table(t, ncol)
    widths = getattr(t, '_tbl_widths', None)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        if widths:
            for ci, cell in enumerate(cells):
                cell.width = Emu(int(widths[ci] * 635))
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ''
            cell = cells[ci]
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Emu(0)
            add_rich(cp, txt, base_bold=(ri == 0))
    doc.add_paragraph()

table_buf = []
while i < len(lines):
    line = lines[i].rstrip()
    i += 1

    # 箇条書きの中に入れ子になったコールアウト（  - > [!warning] …）を平らにする
    mc = re.match(r'^(\s*)[-*]\s*>\s*\[!(\w+)\]\s*(.*)$', line)
    if mc:
        kind = {'warning': '【要注意】', 'important': '【重要】', 'note': '【注】',
                'tip': '【補足】', 'info': '【補足】', 'caution': '【要注意】',
                'danger': '【要注意】'}.get(mc.group(2).lower(), '【注】')
        line = f'{mc.group(1)}- **{kind}** {mc.group(3)}'

    # --- 表 ---
    if line.startswith('|'):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if all(re.fullmatch(r':?-{2,}:?', c) for c in cells if c):
            continue
        table_buf.append(cells); continue
    elif table_buf:
        flush_table(table_buf); table_buf = []

    if not line.strip():
        continue
    if line.startswith('（説教段落') and line.rstrip().endswith('）'):
        continue
    if line.strip() == '---':
        continue

    # --- 見出し ---
    m = re.match(r'^(#{1,4}) (.+)$', line)
    if m:
        lvl, txt = len(m.group(1)), strip_md(m.group(2))
        if lvl == 1:
            continue
        style = {2: 'Heading 1', 3: 'Heading 2', 4: 'Heading 2'}[lvl]
        p = newp(style); p.add_run(txt)
        quote_mode = None
        continue

    # --- 引用ブロックの直前の合図 ---
    if 'ESV（Mark 3:20–35）' in line:
        p = newp(); add_rich(p, line, base_bold=True); quote_mode = 'esv'; continue
    if '新改訳2017（マルコ3:20–35）' in line:
        p = newp(); add_rich(p, line, base_bold=True); quote_mode = 'sky'; continue
    if 'イザヤ49:24–25' in line and '旧約の背景' in line:
        p = newp(); add_rich(p, line); quote_mode = ('ref', 'イザヤ49:'); continue
    if 'マルコ10:29–30 との対応' in line:
        p = newp(); add_rich(p, line); quote_mode = ('ref', 'マルコ10:'); continue

    # --- 引用ブロック本体 ---
    if line.startswith('>'):
        content = line[1:].strip()
        if not content: continue
        p = newp(); pf = p.paragraph_format
        pf.left_indent = INDENT; pf.space_after = SPACE
        if quote_mode == 'esv':
            m2 = re.match(r'^(\d+) (.*)$', content)
            if m2:
                r = p.add_run(f'ESV Mark 3:{m2.group(1)}　{m2.group(2)}'); setfont(r, S.VERSE_PT, color=GREY)
            else:
                r = p.add_run(content); setfont(r, S.VERSE_PT, color=GREY)
        elif quote_mode == 'sky':
            m2 = re.match(r'^(\d+)　(.*)$', content)
            if m2:
                r = p.add_run(f'マルコ3:{m2.group(1)}'); setfont(r, S.VERSE_PT, bold=True, color=BLUE)
                r = p.add_run('　');                 setfont(r, S.VERSE_PT, bold=True, color=BLUE)
                r = p.add_run(m2.group(2));              setfont(r, S.VERSE_PT, bold=True)
            else:
                add_rich(p, content, base_bold=True, size=S.VERSE_PT)
        elif isinstance(quote_mode, tuple):
            book = quote_mode[1]
            m2 = re.match(r'^(\d+)　(.*)$', content)
            if m2:
                r = p.add_run(f'{book}{m2.group(1)}　{strip_md(m2.group(2))}'); setfont(r, S.VERSE_PT, color=ORANGE)
            else:
                r = p.add_run(strip_md(content)); setfont(r, S.VERSE_PT, color=ORANGE)
        else:
            # Obsidian のコールアウト記法（> [!warning] …）をWordで読める形に直す
            m3 = re.match(r'^\[!(\w+)\]\s*(.*)$', content)
            if m3:
                kind = {'warning': '【要注意】', 'important': '【重要】',
                        'note': '【注】', 'tip': '【補足】', 'info': '【補足】',
                        'caution': '【要注意】', 'danger': '【要注意】'}.get(m3.group(1).lower(), '【注】')
                setfont(p.add_run(kind + ' '), bold=True)
                add_rich(p, m3.group(2), base_bold=True)
            else:
                add_rich(p, content)
        continue

    # --- 箇条書き ---
    m = re.match(r'^(\s*)[-*] (.+)$', line)
    if m:
        depth = len(m.group(1)) // 2
        p = newp('List Bullet') if 'List Bullet' in [s.name for s in doc.styles] else newp()
        pf = p.paragraph_format
        pf.left_indent = Emu(228600 * (depth + 1)); pf.space_after = SPACE
        if p.style.name != 'List Bullet':
            add_rich(p, '・' + m.group(2))
        else:
            add_rich(p, m.group(2))
        continue
    m = re.match(r'^(\s*)(\d+)\. (.+)$', line)
    if m:
        p = newp(); pf = p.paragraph_format
        pf.left_indent = Emu(228600); pf.space_after = SPACE
        add_rich(p, f'{m.group(2)}. {m.group(3)}')
        continue

    # --- 通常段落 ---
    quote_mode = None
    p = newp(); p.paragraph_format.space_after = SPACE
    add_rich(p, line)

if table_buf: flush_table(table_buf)
doc.save(OUT)
print('saved:', OUT, '段落数:', len(doc.paragraphs), '表:', len(doc.tables))
