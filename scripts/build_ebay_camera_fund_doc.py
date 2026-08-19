from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "eBay 7月カメラ資金12万円計画 毎日の紙.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(95, 95, 95)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "AAB7C4"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, color=INK, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def paragraph(doc, text="", size=11, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold)
    return p


def heading(doc, text, level=1):
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, BLUE, 14, 7
    else:
        size, color, before, after = 12, DARK_BLUE, 10, 5
    return paragraph(doc, text, size=size, color=color, bold=True, before=before, after=after)


def checkbox_line(doc, text):
    p = paragraph(doc, after=3)
    r = p.add_run("□ ")
    set_font(r, size=11, color=INK, bold=True)
    r = p.add_run(text)
    set_font(r, size=11, color=INK)
    return p


def add_meta_strip(doc):
    table = doc.add_table(rows=2, cols=4)
    set_cell_margins(table)
    set_table_width(table, [1800, 2520, 2520, 2520])
    labels = ["目標", "7月の数字", "出品基準", "レビュー"]
    values = ["手元12万円", "販売総額 約$1,000", "$30利益候補", "毎週火曜"]
    for c, label in enumerate(labels):
        cell = table.cell(0, c)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_font(r, size=9.5, color=DARK_BLUE, bold=True)
    for c, value in enumerate(values):
        cell = table.cell(1, c)
        set_cell_shading(cell, WHITE)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_font(r, size=11, color=INK, bold=True)
    paragraph(doc, "", after=4)


def add_week_table(doc, title, rows, review_text):
    heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=3)
    set_cell_margins(table, top=70, bottom=70, start=100, end=100)
    set_table_width(table, [620, 5960, 2780])
    headers = ["済", "今日やること", "メモ"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=9.5, color=DARK_BLUE, bold=True)
    for item, note in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_borders(cell)
            set_cell_shading(cell, WHITE)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cells[0].paragraphs[0].add_run("□")
        set_font(r, size=12, color=INK, bold=True)
        r = cells[1].paragraphs[0].add_run(item)
        set_font(r, size=10.5, color=INK)
        r = cells[2].paragraphs[0].add_run(note)
        set_font(r, size=9.5, color=MUTED)
    paragraph(doc, review_text, size=10, color=DARK_BLUE, bold=True, before=5, after=4)


def add_daily_table(doc, title, rows):
    heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=4)
    set_cell_margins(table, top=70, bottom=70, start=90, end=90)
    set_table_width(table, [1180, 560, 5540, 2080])
    headers = ["日付", "済", "その日の主作業", "結果メモ"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=9.2, color=DARK_BLUE, bold=True)
    for date, task, note in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_borders(cell)
            set_cell_shading(cell, WHITE)
        for idx, txt in enumerate((date, "□", task, note)):
            p = cells[idx].paragraphs[0]
            if idx in (0, 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            set_font(r, size=9.5 if idx != 2 else 10, color=INK if idx != 3 else MUTED, bold=(idx == 1))


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "eBay 7月カメラ資金12万円計画"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "毎週レビュー: 6/30, 7/7, 7/14, 7/21, 7/28"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.runs[0], size=9, color=MUTED)

    paragraph(doc, "eBay 7月カメラ資金12万円計画", size=22, color=RGBColor(0, 0, 0), bold=True, after=3)
    paragraph(doc, "毎日のやること紙 | 期限: 2026年7月31日 | 目的: Luna Ultra等のVLOGカメラ資金を作る", size=10.5, color=MUTED, after=12)
    add_meta_strip(doc)

    heading(doc, "使い方", 1)
    for line in [
        "毎朝、この紙の今日の日付だけ見る。",
        "終わったらチェックを入れ、結果メモに数字を書く。",
        "週末またはレビュー日に、候補枚数・出品数・売上見込みだけ確認する。",
        "迷ったら、$30利益候補を増やすより先に、出品できる状態を優先する。",
    ]:
        checkbox_line(doc, line)

    heading(doc, "毎週レビューで見る数字", 1)
    table = doc.add_table(rows=1, cols=4)
    set_cell_margins(table)
    set_table_width(table, [2340, 2340, 2340, 2340])
    headers = ["候補", "出品", "反応", "お金"]
    body = [
        "候補枚数\n$30利益候補\nロット候補",
        "出品数\n写真未完\n説明文未完",
        "ウォッチ数\nオファー数\n質問数",
        "販売総額\n手元見込み\n不足額",
    ]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, color=DARK_BLUE, bold=True)
    row = table.add_row().cells
    for i, text in enumerate(body):
        set_cell_borders(row[i])
        p = row[i].paragraphs[0]
        r = p.add_run(text)
        set_font(r, size=10, color=INK)

    doc.add_page_break()

    add_week_table(
        doc,
        "0週目: 6/24〜6/30 | 売れる状態に戻す",
        [
            ("eBayにログインし、出品できる状態か確認する。", "ログイン / 出品 / 制限"),
            ("本人確認・出金口座・販売制限の有無を確認する。", "止まりどころを先に潰す"),
            ("発送方法を1つに絞り、7インチ/LPの送料目安を確認する。", "米国向け優先"),
            ("梱包材、はかり、段ボール、保護材の手持ちを確認する。", "不足分を書く"),
            ("3F=C「とまどい」の状態確認と撮影をする。", "最初の単品候補"),
            ("$30利益候補を探す箱・棚を1か所決める。", "7月に見る場所"),
        ],
        "6/30レビュー: eBayは出品可能か / 送料で詰まっていないか / 100枚を見る場所は決まったか",
    )

    add_daily_table(
        doc,
        "1週目: 7/1〜7/7 | 100枚を見る、40枚候補を抜く",
        [
            ("7/1", "レコード100枚を目視で抜き、通し番号を付ける。", "候補__枚"),
            ("7/2", "100枚をスマホで表裏ざっくり撮る。", "撮影__枚"),
            ("7/3", "型番・レーベル・年が読めるものを優先してメモする。", "型番__枚"),
            ("7/4", "Discogsで上位20枚だけ確認する。", "$30候補__枚"),
            ("7/5", "Discogsで追加20枚を確認し、40枚候補に近づける。", "$30候補__枚"),
            ("7/6", "ロット候補と後回し盤を分ける。", "ロット__件"),
            ("7/7", "3F=Cを1点目として出品する。週レビューをする。", "出品__件"),
        ],
    )

    doc.add_page_break()

    add_daily_table(
        doc,
        "2週目: 7/8〜7/14 | まず20点出品する",
        [
            ("7/8", "単品候補20点を選び、状態を簡単にグレーディングする。", "20点選定"),
            ("7/9", "写真を撮り直す必要がある盤だけ丁寧に撮る。", "写真__点"),
            ("7/10", "出品タイトルと英語説明文テンプレを作る。", "テンプレ完成"),
            ("7/11", "価格を強気固定 + Best Offerで決め、最低受け入れ価格を書く。", "下限記入"),
            ("7/12", "10点出品する。", "出品__件"),
            ("7/13", "さらに10点出品する。", "累計__件"),
            ("7/14", "ロット候補を1件作る。週レビューをする。", "売上見込み__円"),
        ],
    )

    add_daily_table(
        doc,
        "3週目: 7/15〜7/21 | 売上を作る、追加20点出す",
        [
            ("7/15", "反応のある出品を確認し、ウォッチが多いものを調整する。", "ウォッチ__件"),
            ("7/16", "反応のない出品の写真・タイトルを直す。", "修正__件"),
            ("7/17", "追加20点を選ぶ。", "追加候補__点"),
            ("7/18", "追加10点を出品する。", "累計__件"),
            ("7/19", "追加10点を出品する。", "累計__件"),
            ("7/20", "ロット出品を3〜5件作る。", "ロット__件"),
            ("7/21", "売上・手数料・送料・手元見込みを表にする。週レビューをする。", "不足__円"),
        ],
    )

    doc.add_page_break()

    add_daily_table(
        doc,
        "4週目: 7/22〜7/28 | 不足分を埋める",
        [
            ("7/22", "未出品の$30利益候補を追加で出す。", "追加__件"),
            ("7/23", "売れない単品の一部をロットへ切り替える。", "切替__件"),
            ("7/24", "早く売りたい盤だけ価格を少し下げる。宝は安く流しすぎない。", "調整__件"),
            ("7/25", "発送済み案件の追跡・受取状況を確認する。", "追跡OK"),
            ("7/26", "入金予定日を確認する。", "入金予定__円"),
            ("7/27", "カメラ購入可能額を計算する。", "可能額__円"),
            ("7/28", "Luna Ultra / DJIの価格と保証を再確認する。週レビューをする。", "買う/待つ"),
        ],
    )

    add_daily_table(
        doc,
        "最終調整: 7/29〜7/31 | 買う判断",
        [
            ("7/29", "7月売上総額と手元見込み額を確認する。", "総額__円"),
            ("7/30", "未着金分と実キャッシュを分け、カメラ購入予算を決める。", "予算__円"),
            ("7/31", "買う / 待つ / 予算を下げる の判断をする。", "結論"),
        ],
    )

    heading(doc, "一番大事な確認", 1)
    for line in [
        "最初に確認するのは、eBayが売れる状態かどうか。",
        "$30利益候補が見つからない日は、枚数を増やすより基準を見直す。",
        "1件売れたら、販売より先に発送までの手順を記録する。",
        "12万円は目標。未着金と実キャッシュは分けて判断する。",
    ]:
        checkbox_line(doc, line)

    doc.save(OUT)


if __name__ == "__main__":
    build()

